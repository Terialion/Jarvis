from __future__ import annotations
from pathlib import Path
from typing import Any, Dict

DESCRIPTION = "Docx 生成适配入口"
ICON = "📄"


def execute(text: str = "", output: str = "", user_input: str = "", **kwargs) -> Dict[str, Any]:
    try:
      from docx import Document
    except Exception:
      return {"status": "error", "code": "missing_dependency", "message": "python-docx not installed"}

    out = (output or "d:/jarvis/workspace/auto_generated.docx").strip()
    doc = Document()
    doc.add_heading('Generated Document', level=1)
    doc.add_paragraph((text or user_input or 'Hello from minimax-docx adapter')[:4000])
    Path(out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return {"status": "success", "output": out, "notes": "docx_generated"}
